# -*- coding: utf-8 -*-
"""
Formulario de verificación de datos visita - Optimizado y Responsivo
"""

import io
import json
from datetime import datetime

import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    from streamlit_js_eval import get_geolocation
    GEO_OK = True
except Exception:
    GEO_OK = False

# --------------------------------------------------------------------------
# CONFIGURACIÓN GENERAL
# --------------------------------------------------------------------------
st.set_page_config(
    page_title="Formulario - Visita de clientes",
    page_icon="🏦",
    layout="wide",
    initial_sidebar_state="expanded",
)

EXCEL_COLUMNS = [
    "RECNO", "PEPAIS", "PETDOC", "PENDOC", "CODCLI", "BCEMP", "BCSUC", "BCMDA",
    "BCPAP", "BCCTA", "BCOPER", "BCSBOP", "BCTOP", "BCMOD", "CODCRE", "REGION",
    "ZONA", "AGENCIA", "CLIENTE", "DIRECCION_DOM", "DISTRITO_DOM",
    "PROVINCIA_DOM", "DEPARTAMENTO_DOM", "DIRECCION_NEG", "DISTRITO_NEG",
    "PROVINCIA_NEG", "DEPARTAMENTO_NEG", "ACTIVIDAD_ECON", "ANALISTA",
    "PRODUCTO_CAJA", "SALDO_MN", "SALDO_VIGE", "SALDO_REFI", "SALDO_VENC",
    "SALDO_JUDI", "MORA_CONT", "TIPO_SBS", "FECDES", "IMPDESEMB_MN",
    "COD_MODULO", "MODULO", "COD_TIPO_OPERACION", "TIPO_OPERACION",
    "ANALISTA_EVAL", "USUARIO_APROB", "USUARIO_DESEM", "FECHA_EVAL",
    "DIAS_ATRASO", "ESTADO_CREDITO", "ATRANT_1M", "ATRANT_2M", "ATRANT_3M",
    "ATRANT_4M", "ATRANT_5M", "ATRANT_6M", "TIPO_SOLI", "NUMERO_CUOTAS",
    "CUOTAS_PAGADAS", "TIPO", "SEGMENTACION_MYPE", "CATEG_RESULTANTE",
    "CATEG_RESULTANTE_SINALIN", "CUENTA_AVAL", "FECHA_UTLPAGO", "UAI_IND",
    "ESTRATO", "TIPO_EXPEDIENTE",
]

NARANJA = "C8102E"   # Color institucional principal
AZUL = "1B3A5C"      # Color secundario / Encabezados
VERDE = "137333"     # Éxito / OK
ROJO = "A50E0E"      # Riesgo / Alerta

CUSTOM_CSS = f"""
<style>
.stApp {{ background-color: #f8fafc; }}
section[data-testid="stSidebar"] {{ background-color: #{AZUL}; }}
section[data-testid="stSidebar"] * {{ color: #ffffff !important; }}
h1, h2, h3 {{ color: #{AZUL}; font-weight: 700; }}

/* Botones con estilo moderno */
div.stButton > button {{
    background-color: #{AZUL}; color: white; border: none;
    border-radius: 6px; font-weight: 600; padding: 0.5rem 1rem;
    transition: all 0.3s ease;
}}
div.stButton > button:hover {{ background-color: #{NARANJA}; color: white; }}

/* Tarjetas contenedoras */
.card {{
    background: white; padding: 1.5rem; border-radius: 12px;
    border: 1px solid #e2e8f0;
    box-shadow: 0 4px 6px -1px rgba(0,0,0,0.05); margin-bottom: 1.2rem;
}}

/* Estilos de la Matriz de Riesgos */
.validation-box {{
    border: 1px solid #cbd5e1; border-radius: 12px; padding: 1.5rem;
    background: #ffffff; margin: 1rem 0;
}}
.validation-title {{
    font-size: 1.2rem; font-weight: 700; color: #{AZUL}; margin-bottom: 1rem;
    display: flex; align-items: center; gap: 0.5rem;
}}
.macro-header-stripe {{
    background-color: #{AZUL}; color: white; padding: 10px 14px;
    font-weight: 600; border-radius: 6px; margin-top: 14px;
    margin-bottom: 8px; font-size: 0.95rem; box-shadow: inset 0 -2px 0 rgba(0,0,0,0.15);
}}
.criterion-row-active {{
    background-color: #fef2f2; border-left: 4px solid #{ROJO};
    padding: 6px 10px; border-radius: 0 6px 6px 0; margin-bottom: 4px;
}}
.criterion-row-inactive {{
    background-color: #f8fafc; border-left: 4px solid #cbd5e1;
    padding: 6px 10px; border-radius: 0 6px 6px 0; margin-bottom: 4px;
}}

.badge-ok {{ background:#e6f4ea; color:#137333; padding:4px 10px; border-radius:12px; font-size:0.8rem; font-weight:600; }}
.badge-pend {{ background:#fce8e6; color:#a50e0e; padding:4px 10px; border-radius:12px; font-size:0.8rem; font-weight:600; }}

/* Ajustes de adaptabilidad móvil avanzada */
@media (max-width: 768px) {{
    .card {{ padding: 1rem; }}
    h1 {{ font-size: 1.6rem !important; }}
    h2 {{ font-size: 1.3rem !important; }}
    .macro-header-stripe {{ font-size: 0.85rem; padding: 8px 10px; }}
    /* Forzar que las columnas de Streamlit se comporten de forma fluida en pantallas pequeñas */
    [data-testid="column"] {{ width: 100% !important; flex: 1 1 100% !important; }}
}}
</style>
"""
st.markdown(CUSTOM_CSS, unsafe_allow_html=True)


# --------------------------------------------------------------------------
# HELPERS
# --------------------------------------------------------------------------
def safe_str(v, default=""):
    if v is None:
        return default
    try:
        if pd.isna(v):
            return default
    except Exception:
        pass
    s = str(v).strip()
    return default if s.lower() in ("nan", "none") else s


def safe_float(v, default=0.0):
    try:
        f = float(v)
        if pd.isna(f):
            return default
        return f
    except Exception:
        return default


def fmt_money(v):
    return f"S/. {safe_float(v):,.2f}"


def init_state():
    defaults = {
        "clientes_df": None,
        "cliente_actual": {},
        "visitas": {},   
        "garantias": [],
        "rcc": [],
        "validaciones_marcadas": {},  
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


init_state()


def buscar_cliente_por_dni(dni_input, df):
    if not dni_input or df is None or len(df) == 0:
        return None
    mask = (df.get("PENDOC", pd.Series("", index=df.index)).astype(str).str.strip() == str(dni_input).strip())
    resultados = df[mask]
    if len(resultados) > 0:
        return resultados.iloc[0].to_dict()
    return None


# --------------------------------------------------------------------------
# VALIDACIÓN DE CRITERIOS AUTOMÁTICOS
# --------------------------------------------------------------------------
def validar_visita():
    validaciones = {
        "documentos_enmiendas": False,
        "documentos_inconsistentes": False,
        "documentos_sin_datos": False,
        "documentos_sin_firmas": False,
        "documentos_duplicados": False,
        "sin_sustento_actividad": False,
        "sin_sustento_ingresos": False,
        "sin_sustento_activos": False,
        "conyuge_omitido": False,
        "credito_reprogramado": False,
        "credito_refinanciado": False,
        "calificacion_diferente": False,
    }
    
    cliente = st.session_state.cliente_actual
    if not cliente:
        return validaciones

    if not safe_str(cliente.get("CLIENTE")):
        validaciones["documentos_sin_datos"] = True
    
    if safe_str(cliente.get("DIAS_ATRASO")) and int(safe_float(cliente.get("DIAS_ATRASO"))) > 0:
        validaciones["calificacion_diferente"] = True
    
    visitas = st.session_state.visitas
    for clave in ["domicilio", "negocio", "aval"]:
        if clave not in visitas:
            validaciones["sin_sustento_actividad"] = True
            break
        visita = visitas[clave]
        if not visita.get("foto_bytes"):
            validaciones["documentos_sin_firmas"] = True
    
    return validaciones


# --------------------------------------------------------------------------
# FUNCIÓN PANEL DE VALIDACIÓN (MATRIZ IDENTICA A LA CAPTURA)
# --------------------------------------------------------------------------
def mostrar_panel_validacion():
    st.markdown('<div class="validation-box">', unsafe_allow_html=True)
    st.markdown('<div class="validation-title">📋 Matriz de Control - Criterios para Visita a Clientes</div>', unsafe_allow_html=True)
    
    validaciones_auto = validar_visita()
    
    # Estructura jerárquica de la tabla adjunta en la imagen
    matriz_criterios = {
        "Indicio de dolo o fraude en la evaluación de créditos": {
            "documentos_enmiendas": "Documentos con enmiendas",
            "documentos_inconsistentes": "Documentos con datos inconsistentes",
            "documentos_sin_datos": "Documentos sin datos del cliente",
            "documentos_sin_firmas": "Documentos sin firmas o que no coinciden",
            "documentos_duplicados": "Documentos duplicados en más de un cliente"
        },
        "Evaluaciones deficientes o con sustento insuficiente": {
            "sin_sustento_actividad": "No se evidenció sustento de actividad económica",
            "sin_sustento_ingresos": "No se evidenció sustento de ingresos",
            "sin_sustento_activos": "No se evidenció sustento de activos representativos",
            "conyuge_omitido": "Se omitió al cónyuge"
        },
        "Créditos reprogramados y refinanciados": {
            "credito_reprogramado": "Reprogramado",
            "credito_refinanciado": "Refinanciado"
        },
        "Clientes con créditos con calificación diferente a normal a la fecha de revisión": {
            "calificacion_diferente": "Indicar la calificación a la fecha de revisión"
        }
    }
    
    for macro_cat, sub_criterios in matriz_criterios.items():
        st.markdown(f'<div class="macro-header-stripe">{macro_cat}</div>', unsafe_allow_html=True)
        
        for key, label in sub_criterios.items():
            # Mantiene el estado cruzado entre cálculo automático e interacción manual
            is_checked = st.session_state.validaciones_marcadas.get(key, validaciones_auto.get(key, False))
            
            # Contenedor responsivo nativo de Streamlit
            col_check, col_text = st.columns([0.08, 0.92])
            with col_check:
                sub_check = st.checkbox(
                    label="",
                    value=is_checked,
                    key=f"chk_matriz_{key}",
                    label_visibility="collapsed"
                )
                st.session_state.validaciones_marcadas[key] = sub_check
                
            with col_text:
                if sub_check:
                    st.markdown(f'<div class="criterion-row-active"><span style="color:#A50E0E; font-weight:700;">[X] {label}</span></div>', unsafe_allow_html=True)
                else:
                    st.markdown(f'<div class="criterion-row-inactive"><span style="color:#475569;">[ ] {label}</span></div>', unsafe_allow_html=True)
    
    st.divider()
    total_marcados = sum(1 for v in st.session_state.validaciones_marcadas.values() if v)
    if total_marcados == 0:
        st.success("✅ Visita conforme: Ningún criterio de riesgo alertado.")
    else:
        st.error(f"⚠️ Atención: Se han seleccionado/detectado {total_marcados} criterios de riesgo en esta evaluación.")
    st.markdown('</div>', unsafe_allow_html=True)


# --------------------------------------------------------------------------
# SIDEBAR: CARGA DE ARCHIVOS Y CONTROL
# --------------------------------------------------------------------------
with st.sidebar:
    st.markdown("### 📂 Base de datos")
    excel_file = st.file_uploader(
        "Cargar archivo .xlsx",
        type=["xlsx", "xls"],
        help="Debe contener las columnas de control institucional.",
    )
    if excel_file is not None:
        try:
            df_upload = pd.read_excel(excel_file, dtype=str)
            df_upload.columns = [c.strip().upper() for c in df_upload.columns]
            st.session_state.clientes_df = df_upload
            st.success(f"✅ {len(df_upload)} registros listos")
        except Exception as e:
            st.error(f"Error al procesar archivo: {e}")

    st.divider()
    df = st.session_state.clientes_df
    if df is not None:
        st.markdown("### 🔍 Buscador de Clientes")
        busq = st.text_input("Buscar por Nombre o Código", key="busqueda_sidebar")
        if busq:
            b = busq.strip().lower()
            mask = (
                df.get("CODCLI", pd.Series("", index=df.index)).astype(str).str.contains(b, case=False, na=False) |
                df.get("CLIENTE", pd.Series("", index=df.index)).astype(str).str.contains(b, case=False, na=False)
            )
            resultados = df[mask]
        else:
            resultados = df

        if len(resultados) > 0:
            opciones = resultados.apply(
                lambda r: f"{safe_str(r.get('CODCLI'))} | {safe_str(r.get('CLIENTE'))}", axis=1
            ).tolist()
            sel = st.selectbox("Resultados:", opciones)
            if sel and st.button("📥 Cargar Datos Seleccionados", use_container_width=True):
                idx_sel = opciones.index(sel)
                st.session_state.cliente_actual = resultados.iloc[idx_sel].to_dict()
                st.session_state.visitas = {}
                st.session_state.validaciones_marcadas = {}
                st.rerun()

    st.divider()
    if st.session_state.cliente_actual:
        st.warning(f"Trabajando con:\n**{safe_str(st.session_state.cliente_actual.get('CLIENTE'))}**")
        if st.button("🗑️ Resetear Formulario Actual", use_container_width=True):
            st.session_state.cliente_actual = {}
            st.session_state.visitas = {}
            st.session_state.garantias = []
            st.session_state.rcc = []
            st.session_state.validaciones_marcadas = {}
            st.rerun()

# --------------------------------------------------------------------------
# CALLBACK PARA AUTOLLENADO DIRECTO DESDE EL INPUT DNI
# --------------------------------------------------------------------------
def cambiar_dni_callback():
    dni_digitado = st.session_state.dni_interactivo.strip()
    if dni_digitado and st.session_state.clientes_df is not None:
        cliente_encontrado = buscar_cliente_por_dni(dni_digitado, st.session_state.clientes_df)
        if cliente_encontrado:
            st.session_state.cliente_actual = cliente_encontrado
            st.session_state.visitas = {}
            st.session_state.validaciones_marcadas = {}
        else:
            st.sidebar.error("DNI no localizado en la base de datos actual.")


cliente = st.session_state.cliente_actual

st.title("Visita a Clientes")
st.caption("Formulario digital de verificación institucional")

tabs = st.tabs([
    "1️⃣ Cliente y Crédito",
    "2️⃣ Historial y Riesgo",
    "3️⃣ Visita Domicilio",
    "4️⃣ Visita Negocio",
    "5️⃣ Ingresos y Gastos",
    "6️⃣ Garantías y Aval",
    "7️⃣ Generar Reporte",
])

# --------------------------------------------------------------------------
# TAB 1 — DATOS DEL CLIENTE Y CRÉDITO (BÚSQUEDA REACTIVA POR DNI)
# --------------------------------------------------------------------------
with tabs[0]:
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("Datos Generales del Titular")
    
    # Input DNI con callback inmediato de autollenado
    st.text_input(
        "🔍 Ingrese DNI del Titular (Presione Enter para Autofoco/Autollenado)", 
        value=safe_str(cliente.get("PENDOC")),
        key="dni_interactivo",
        on_change=cambiar_dni_callback
    )
    
    c1, c2, c3 = st.columns([1, 1, 1])
    with c1:
        agencia = st.text_input("Agencia", value=safe_str(cliente.get("AGENCIA")))
        dni = st.text_input("DNI / LE Documento", value=safe_str(cliente.get("PENDOC")), disabled=True)
        codcli = st.text_input("Código de cliente", value=safe_str(cliente.get("CODCLI")))
    with c2:
        titular = st.text_input("Nombre del titular", value=safe_str(cliente.get("CLIENTE")))
        cuenta = st.text_input("Cuenta cliente", value=safe_str(cliente.get("BCCTA")))
        operacion = st.text_input("Nro. de operación", value=safe_str(cliente.get("BCOPER")))
    with c3:
        analista = st.text_input("Analista vigente", value=safe_str(cliente.get("ANALISTA")))
        analista_eval = st.text_input("Analista evaluador", value=safe_str(cliente.get("ANALISTA_EVAL")))
        aprobado_por = st.text_input("Aprobado por", value=safe_str(cliente.get("USUARIO_APROB")))
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("Datos del crédito")
    c1, c2, c3 = st.columns([1, 1, 1])
    with c1:
        importe = st.number_input("Importe desembolsado (S/.)", value=safe_float(cliente.get("IMPDESEMB_MN")), format="%.2f")
        saldo_capital = st.number_input("Saldo capital (S/.)", value=safe_float(cliente.get("SALDO_MN")), format="%.2f")
        tipo_credito = st.text_input("Tipo de crédito", value=safe_str(cliente.get("PRODUCTO_CAJA")))
    with c2:
        tipo_sbs = st.text_input("Tipo según SBS", value=safe_str(cliente.get("TIPO_SBS")))
        fecha_desembolso = st.text_input("Fecha de desembolso", value=safe_str(cliente.get("FECDES")))
        cuotas_pagadas = st.text_input("Nro. cuotas pagadas", value=safe_str(cliente.get("CUOTAS_PAGADAS")))
    with c3:
        dias_atraso = st.text_input("Días de atraso", value=safe_str(cliente.get("DIAS_ATRASO")))
        prom_mora = st.text_input("Promedio de mora", value=safe_str(cliente.get("MORA_CONT")))
        calificacion = st.text_input("Calificación", value=safe_str(cliente.get("CATEG_RESULTANTE")))
        
    c1, c2, c3 = st.columns([1, 1, 1])
    with c1:
        rubro = st.text_input("Rubro / Actividad económica", value=safe_str(cliente.get("ACTIVIDAD_ECON")))
    with c2:
        sector = st.text_input("Segmentación MYPE", value=safe_str(cliente.get("SEGMENTACION_MYPE")))
    with c3:
        modulo = st.text_input("Módulo", value=safe_str(cliente.get("MODULO")))
    st.markdown("</div>", unsafe_allow_html=True)
    
    # Inserción estructurada de la Matriz de Control de Riesgos
    mostrar_panel_validacion()


# --------------------------------------------------------------------------
# TAB 2 — HISTORIAL CREDITICIO
# --------------------------------------------------------------------------
with tabs[1]:
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("3. Historial crediticio interno")
    df = st.session_state.clientes_df
    if df is not None and (dni or codcli):
        mask = pd.Series(False, index=df.index)
        if dni:
            mask = mask | (df.get("PENDOC", "").astype(str) == dni)
        if codcli:
            mask = mask | (df.get("CODCLI", "").astype(str) == codcli)
        hist = df[mask]
        cols_show = [c for c in [
            "AGENCIA", "CODCRE", "ESTADO_CREDITO", "FECDES", "FECHA_UTLPAGO",
            "PRODUCTO_CAJA", "SALDO_MN", "DIAS_ATRASO", "ANALISTA",
        ] if c in hist.columns]
        if len(hist) > 0:
            st.dataframe(hist[cols_show], use_container_width=True, hide_index=True)
        else:
            st.info("No se encontraron registros alternos para el cliente.")
    else:
        st.info("Cargue la base de datos Excel para activar el tracking histórico.")
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("4. Riesgo de sobreendeudamiento")
    deuda_directa_auto = 0.0
    if df is not None and (dni or codcli) and 'hist' in locals() and len(hist) > 0:
        for col in ["SALDO_VIGE", "SALDO_REFI"]:
            if col in hist.columns:
                deuda_directa_auto += pd.to_numeric(hist[col], errors="coerce").fillna(0).sum()

    c1, c2 = st.columns(2)
    with c1:
        deuda_directa = st.number_input("a) Deuda directa (S/.)", value=float(deuda_directa_auto), format="%.2f")
        deuda_potencial = st.number_input("b) Deuda potencial (S/.)", value=0.0, format="%.2f")
        deuda_total = st.number_input("c) Deuda total (S/.)", value=float(deuda_directa) + float(deuda_potencial), format="%.2f")
    with c2:
        resultado_neto_rs = st.number_input("e) Cuota / Resultado neto (%)", value=0.0, format="%.2f")
        pasivo_patrimonio = st.number_input("f) Pasivo / Patrimonio (%)", value=0.0, format="%.2f")
        tipo_deuda = st.selectbox("d) Tipo Moneda", ["Sin identificación", "MN", "ME"])
    st.markdown("</div>", unsafe_allow_html=True)


# --------------------------------------------------------------------------
# BLOQUE MULTIDISPOSITIVO REUTILIZABLE PARA VISITAS
# --------------------------------------------------------------------------
def bloque_verificacion(clave, etiqueta):
    st.markdown("##### 📍 Registro de verificación de campo")
    visitas = st.session_state.visitas
    data = visitas.get(clave, {})

    colf1, colf2 = st.columns(2)
    with colf1:
        fecha_v = st.date_input("Fecha de visita", value=datetime.now().date(), key=f"fecha_{clave}")
    with colf2:
        hora_v = st.time_input("Hora de visita", value=datetime.now().time(), key=f"hora_{clave}")

    entrevista_con = st.text_input("Entrevistado / Contacto:", key=f"entrevista_{clave}")
    comentarios = st.text_area("Notas u Observaciones encontradas:", key=f"comentarios_{clave}")

    st.markdown("**Ubicación Referencial GPS**")
    cgps1, cgps2 = st.columns([1, 2])
    with cgps1:
        capturar = st.button("📡 Geolocalizar Dispositivo", key=f"btn_gps_{clave}", use_container_width=True)
    lat, lon, precision = data.get("lat"), data.get("lon"), data.get("precision")
    if capturar:
        if GEO_OK:
            loc = get_geolocation(key=f"geo_{clave}_{datetime.now().timestamp()}")
            if loc and "coords" in loc:
                lat = loc["coords"]["latitude"]
                lon = loc["coords"]["longitude"]
                precision = loc["coords"].get("accuracy")
            else:
                st.warning("Active los permisos de geolocalización en su navegador móvil/PC.")
        else:
            st.warning("Geolocalización nativa no disponible en este Host.")
            
    with cgps2:
        if lat and lon:
            st.success(f"Ubicación fijada con precisión de (±{precision:.0f} m)")
            st.map(pd.DataFrame({"lat": [lat], "lon": [lon]}), zoom=14, height=160)

    st.markdown("**Panel Fotográfico**")
    cfoto1, cfoto2 = st.columns(2)
    with cfoto1:
        foto_camara = st.camera_input("Activar Cámara Integrada", key=f"camara_{clave}")
    with cfoto2:
        foto_archivo = st.file_uploader("Cargar desde Carrete/Archivos", type=["jpg", "jpeg", "png"], key=f"upload_{clave}")
    foto_final = foto_camara if foto_camara is not None else foto_archivo

    if st.button(f"💾 Confirmar Cierre de Visita: {etiqueta}", key=f"guardar_{clave}", use_container_width=True):
        st.session_state.visitas[clave] = {
            "fecha": str(fecha_v),
            "hora": str(hora_v),
            "entrevista_con": entrevista_con,
            "comentarios": comentarios,
            "lat": lat,
            "lon": lon,
            "precision": precision,
            "foto_bytes": foto_final.getvalue() if foto_final is not None else None,
        }
        st.success(f"✅ Los datos de la visita ({etiqueta}) han sido retenidos temporalmente.")


# --------------------------------------------------------------------------
# TABS DE REGISTROS DE VISITAS INDIVIDUALES
# --------------------------------------------------------------------------
with tabs[2]:
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("5. Dirección del domicilio declarado")
    c1, c2 = st.columns(2)
    with c1:
        direccion_dom = st.text_input("Dirección", value=safe_str(cliente.get("DIRECCION_DOM")), key="dir_dom")
        distrito_dom = st.text_input("Distrito", value=safe_str(cliente.get("DISTRITO_DOM")), key="dist_dom")
    with c2:
        provincia_dom = st.text_input("Provincia", value=safe_str(cliente.get("PROVINCIA_DOM")), key="prov_dom")
        departamento_dom = st.text_input("Departamento", value=safe_str(cliente.get("DEPARTAMENTO_DOM")), key="depto_dom")
    referencia_dom = st.text_area("Referencia urbana de acceso:", key="ref_dom")
    tipo_vivienda = st.selectbox("Condición de tenencia de vivienda", ["Propia", "Familiar", "Alquilada", "Otro"], key="tipo_viv")
    st.divider()
    bloque_verificacion("domicilio", "Domicilio")
    st.markdown("</div>", unsafe_allow_html=True)

with tabs[3]:
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("6. Dirección del negocio / Local comercial")
    c1, c2 = st.columns(2)
    with c1:
        direccion_neg = st.text_input("Dirección Comercial", value=safe_str(cliente.get("DIRECCION_NEG")), key="dir_neg")
        distrito_neg = st.text_input("Distrito", value=safe_str(cliente.get("DISTRITO_NEG")), key="dist_neg")
    with c2:
        provincia_neg = st.text_input("Provincia", value=safe_str(cliente.get("PROVINCIA_NEG")), key="prov_neg")
        departamento_neg = st.text_input("Departamento", value=safe_str(cliente.get("DEPARTAMENTO_NEG")), key="depto_neg")
    referencia_neg = st.text_area("Referencia comercial / Rótulo:", key="ref_neg")
    tipo_negocio = st.text_input("Giro comercial verificado", value=safe_str(cliente.get("ACTIVIDAD_ECON")), key="tipo_neg")
    st.divider()
    bloque_verificacion("negocio", "Negocio")
    st.markdown("</div>", unsafe_allow_html=True)

with tabs[4]:
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("Evaluación Económica: Ventas")
    c1, c2 = st.columns(2)
    with c1:
        actividad_principal = st.text_input("Giro comercial Principal", value=safe_str(cliente.get("ACTIVIDAD_ECON")), key="act_princ")
    with c2:
        otras_actividades = st.text_input("Giro comercial secundario", key="otras_act")
    ventas = st.number_input("Ventas brutas mensuales declaradas (S/.)", value=0.0, format="%.2f", key="ventas")

    st.subheader("Estructura de Margen y Costos")
    c1, c2, c3 = st.columns(3)
    with c1:
        costo_ventas = st.number_input("Costo Directo de Ventas", value=0.0, format="%.2f", key="costo_ventas")
        gastos_admin = st.number_input("Gastos Op. / Administrativos", value=0.0, format="%.2f", key="gastos_admin")
    with c2:
        gastos_financieros = st.number_input("Carga Financiera Vigente", value=0.0, format="%.2f", key="gastos_fin")
        gastos_familiares = st.number_input("Carga de Canasta Familiar", value=0.0, format="%.2f", key="gastos_fam")
    with c3:
        otros_ingresos = st.number_input("Otros Ingresos Declarados", value=0.0, format="%.2f", key="otros_ing")
        caja_bancos = st.number_input("Disponibilidad Caja/Bancos", value=0.0, format="%.2f", key="caja_bancos")

    resultado_neto = ventas + otros_ingresos - costo_ventas - gastos_admin - gastos_financieros - gastos_familiares
    utilidad_neta = resultado_neto - gastos_familiares
    c1, c2 = st.columns(2)
    c1.metric("Resultado Neto de Operación", fmt_money(resultado_neto))
    c2.metric("Excedente / Utilidad Neta Liquida", fmt_money(utilidad_neta))
    st.markdown("</div>", unsafe_allow_html=True)

with tabs[5]:
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("7. Estructura de Garantías Propuestas")
    with st.form("form_garantia", clear_on_submit=True):
        c1, c2, c3 = st.columns(3)
        with c1:
            g_desc = st.text_input("Descripción de la Garantía")
            g_moneda = st.selectbox("Divisa", ["S/.", "US$"])
        with c2:
            g_importe = st.number_input("Valor Comercial Tasado", value=0.0, format="%.2f")
            g_perito = st.text_input("Perito Registrado REPEV")
        with c3:
            g_fecha_tasacion = st.date_input("Fecha última tasación", value=None)
            g_fecha_decl = st.date_input("Fecha de declaración jurada", value=datetime.now().date())
        if st.form_submit_button("➕ Vincular Garantía al Expediente"):
            st.session_state.garantias.append({
                "descripcion": g_desc, "moneda": g_moneda, "importe": g_importe,
                "perito": g_perito, "fecha_tasacion": str(g_fecha_tasacion),
                "fecha_declaracion": str(g_fecha_decl),
            })
    if st.session_state.garantias:
        st.dataframe(pd.DataFrame(st.session_state.garantias), use_container_width=True, hide_index=True)
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("VI. Verificación in situ de Avales / Obligados Solidarios")
    cuenta_aval = st.text_input("Código o Cuenta del Aval", value=safe_str(cliente.get("CUENTA_AVAL")), key="cuenta_aval")
    bloque_verificacion("aval", "Aval")
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("VII. Carga Financiera Consolidada Externa (RCC)")
    with st.form("form_rcc", clear_on_submit=True):
        c1, c2, c3 = st.columns(3)
        with c1:
            r_entidad = st.text_input("IFIs / Entidad acreedora")
        with c2:
            r_rubro = st.text_input("Tipo Crédito / Producto")
        with c3:
            r_saldo = st.number_input("Saldo de Deuda Actual (S/.)", value=0.0, format="%.2f")
        if st.form_submit_button("➕ Indexar Deuda Externa"):
            st.session_state.rcc.append({"entidad": r_entidad, "rubro": r_rubro, "saldo": r_saldo})
    if st.session_state.rcc:
        st.dataframe(pd.DataFrame(st.session_state.rcc), use_container_width=True, hide_index=True)
    st.markdown("</div>", unsafe_allow_html=True)


# --------------------------------------------------------------------------
# LÓGICA DE EXPORTACIÓN FINAL A WORD (.DOCX)
# --------------------------------------------------------------------------
def add_heading(doc, text, size=13, color=AZUL):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_kv_table(doc, pairs, cols=2):
    table = doc.add_table(rows=0, cols=cols * 2)
    table.style = "Light Grid Accent 1"
    row = None
    for i, (k, v) in enumerate(pairs):
        if i % cols == 0:
            row = table.add_row().cells
        c = (i % cols) * 2
        row[c].text = str(k)
        row[c + 1].text = str(v) if v not in (None, "") else "-"
    return table


def visita_a_texto(visitas, clave):
    d = visitas.get(clave)
    if not d:
        return None
    gps = f"{d['lat']:.6f}, {d['lon']:.6f}" if d.get("lat") and d.get("lon") else "No capturada"
    return [
        ("Fecha", d.get("fecha", "-")),
        ("Hora", d.get("hora", "-")),
        ("Entrevista con", d.get("entrevista_con", "-")),
        ("Ubicación GPS", gps),
        ("Comentarios", d.get("comentarios", "-")),
    ]


def generar_reporte():
    doc = Document()
    doc.add_heading("VISITA A CLIENTES DE PEQUEÑA EMPRESA", level=0)
    p = doc.add_paragraph("")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Fecha de reporte: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

    add_heading(doc, "I. Datos del cliente")
    add_kv_table(doc, [
        ("Agencia", agencia), ("DNI/LE Titular", dni),
        ("Titular", titular), ("Cuenta cliente", cuenta),
        ("Analista vigente", analista), ("Analista evaluador", analista_eval),
        ("Importe", fmt_money(importe)), ("Saldo capital", fmt_money(saldo_capital)),
        ("Tipo de crédito", tipo_credito), ("Tipo SBS", tipo_sbs),
        ("Días de atraso", dias_atraso), ("Calificación", calificacion),
        ("Rubro", rubro), ("Sector", sector),
    ])

    add_heading(doc, "II. Riesgo de sobreendeudamiento")
    add_kv_table(doc, [
        ("Deuda directa", fmt_money(deuda_directa)),
        ("Deuda potencial", fmt_money(deuda_potencial)),
        ("Deuda total", fmt_money(deuda_total)),
        ("Cuota/Resultado neto", f"{resultado_neto_rs:.2f}%"),
        ("Pasivo/Patrimonio", f"{pasivo_patrimonio:.2f}%"),
    ])

    for clave, titulo, direccion_info in [
        ("domicilio", "III. Visita al domicilio", [
            ("Dirección", direccion_dom), ("Distrito", distrito_dom),
            ("Provincia", provincia_dom), ("Departamento", departamento_dom),
            ("Tipo de vivienda", tipo_vivienda),
        ]),
        ("negocio", "IV. Visita al negocio", [
            ("Dirección", direccion_neg), ("Distrito", distrito_neg),
            ("Provincia", provincia_neg), ("Departamento", departamento_neg),
            ("Actividad", tipo_negocio),
        ]),
        ("aval", "V. Visita al aval", [("Cuenta aval", cuenta_aval)]),
    ]:
        add_heading(doc, titulo)
        add_kv_table(doc, direccion_info)
        verif = visita_a_texto(st.session_state.visitas, clave)
        if verif:
            doc.add_paragraph("Registro de verificación:").bold = True
            add_kv_table(doc, verif)
            foto_bytes = st.session_state.visitas[clave].get("foto_bytes")
            if foto_bytes:
                doc.add_picture(io.BytesIO(foto_bytes), width=Cm(8))
        else:
            doc.add_paragraph("⚠ No se registró visita de verificación para esta sección.")

    add_heading(doc, "VI. Verificación de ingresos y gastos")
    add_kv_table(doc, [
        ("Ventas", fmt_money(ventas)), ("Otros ingresos", fmt_money(otros_ingresos)),
        ("Costo de ventas", fmt_money(costo_ventas)), ("Gastos administrativos", fmt_money(gastos_admin)),
        ("Gastos financieros", fmt_money(gastos_financieros)), ("Gastos familiares", fmt_money(gastos_familiares)),
        ("Resultado neto", fmt_money(resultado_neto)), ("Utilidad neta", fmt_money(utilidad_neta)),
    ])

    if st.session_state.garantias:
        add_heading(doc, "VII. Garantías")
        for g in st.session_state.garantias:
            add_kv_table(doc, list(g.items()))

    if st.session_state.rcc:
        add_heading(doc, "VIII. Deuda RCC")
        for r in st.session_state.rcc:
            add_kv_table(doc, list(r.items()))

    add_heading(doc, "IX. Validaciones de Riesgos Identificados")
    validaciones_marcadas = [k for k, v in st.session_state.validaciones_marcadas.items() if v]
    if validaciones_marcadas:
        criterios_labels = {
            "documentos_enmiendas": "Documentos con enmiendas",
            "documentos_inconsistentes": "Documentos con datos inconsistentes",
            "documentos_sin_datos": "Documentos sin datos del cliente",
            "documentos_sin_firmas": "Documentos sin firmas o que no coinciden",
            "documentos_duplicados": "Documentos duplicados en más de un cliente",
            "sin_sustento_actividad": "No se evidenció sustento de actividad económica",
            "sin_sustento_ingresos": "No se evidenció sustento de ingresos",
            "sin_sustento_activos": "No se evidenció sustento de activos representativos",
            "conyuge_omitido": "Se omitió al cónyuge",
            "credito_reprogramado": "Reprogramado",
            "credito_refinanciado": "Refinanciado",
            "calificacion_diferente": "Indicar la calificación a la fecha de revisión",
        }
        for key in validaciones_marcadas:
            doc.add_paragraph(f"• {criterios_labels.get(key, key)}", style='List Bullet')
    else:
        doc.add_paragraph("Sin observaciones críticas identificadas en la matriz de control.")

    add_heading(doc, "Firmas y Conformidad")
    add_kv_table(doc, [
        ("Evaluado por (Firma)", ""), ("Fecha de Cierre", datetime.now().strftime("%d/%m/%Y")),
        ("Revisado por (Riesgos)", ""), ("Fecha de Conformidad", ""),
    ])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# --------------------------------------------------------------------------
# TAB 7 — PANEL DE DESCARGAS Y RESUMEN GENERAL
# --------------------------------------------------------------------------
with tabs[6]:
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("Estado de las inspecciones mandatorias")
    for clave, etiqueta in [("domicilio", "Domicilio"), ("negocio", "Negocio"), ("aval", "Aval")]:
        if clave in st.session_state.visitas:
            d = st.session_state.visitas[clave]
            st.markdown(f"**{etiqueta}** — <span class='badge-ok'>Cerrada Exitosamente</span>", unsafe_allow_html=True)
            st.caption(f"{d['fecha']} {d['hora']} · {d.get('entrevista_con') or 'Sin Informante Directo'}")
        else:
            st.markdown(f"**{etiqueta}** — <span class='badge-pend'>Pendiente de Inspección</span>", unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("Generación de Documentación Oficial")
    st.caption("Consolida el archivo ejecutable estructurado con fotos, matrices y geolocalizaciones incorporadas.")
    if st.button("Construir Reporte Corporativo (.docx)", type="primary"):
        buf = generar_reporte()
        nombre = f"Visita_{safe_str(titular, 'cliente').replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        st.download_button(
            "⬇️ Descargar Reporte Generado",
            data=buf,
            file_name=nombre,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
        st.success("Documento Word listo para su almacenamiento.")
    st.markdown("</div>", unsafe_allow_html=True)
